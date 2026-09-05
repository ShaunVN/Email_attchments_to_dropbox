"""
Generate a Word (.docx) document with a line-by-line explanation of the main script.
This script uses python-docx to produce the file: HSBC_script_documentation.docx
"""
from docx import Document

DESCRIPTION = r"""
Email attachment automation for HSBC monthly statements.

This document contains a top-to-bottom description of the main script, its behavior, and important notes.

(Condensed for readability; full explanation follows in sections.)

1. Purpose: find Gmail messages from HSBC with subject "Your Monthly HSBC Bank Statement", extract PDF attachments, skip duplicates, save locally (to your Dropbox-sync folder) by default, or upload to Dropbox only when explicitly enabled and correctly scoped.

2. Key pieces: imaplib for Gmail, email package for parsing, dropbox SDK optional, dotenv for local .env config, processed_statements.json for checkpointing.

3. Notable new features in the updated script:
   - Persistent processed_hashes stored in processed_statements.json to avoid duplicates across runs
   - --dry-run flag to preview actions without writing files, moving emails, or updating state
   - Inline comments added throughout the file for maintainability
   - Introduction of UPLOAD_TO_DROPBOX env flag (defaults to false). By default the script uses LOCAL_DEST_FOLDER and does not upload.
   - State file has two keys: last_processed_date (ISO timestamp) and processed_hashes (list of sha256 hex digests)

4. How to use:
  - Install dependencies: python -m pip install -r requirements.txt
  - Prepare .env in repo root with your credentials (do NOT commit it)
  - Ensure LOCAL_DEST_FOLDER is set to a Dropbox-synced local path where attachments should be saved
  - By default: Run preview: python email_attachments_to_dropbox.py --dry-run (will not save/move)
  - Live local-save: python email_attachments_to_dropbox.py (saves files into LOCAL_DEST_FOLDER)
  - To enable Dropbox uploads: set UPLOAD_TO_DROPBOX=true and provide a DROPBOX_ACCESS_TOKEN with appropriate scopes (files.content.write). Prefer using Full Dropbox scope if you want uploads outside the App sandbox

5. Important caveats & tips
  - Gmail may require an app-specific password and IMAP enabled
  - IMAP label handling varies; if automatic label creation fails, create the "HomeLoans" label manually in Gmail
  - The script stores processed hashes and last processed date in processed_statements.json in the repo root; keep this file private
  - By default uploads are disabled to avoid unexpected App-folder sandboxing. If you enable uploads, confirm the Dropbox app permission model and token scopes before running live

"""

def make_doc():
    doc = Document()
    doc.add_heading('HSBC Email Attachment Automation - Script Documentation', level=1)
    doc.add_paragraph('Generated documentation for email_attachments_to_dropbox.py')
    for para in DESCRIPTION.split('\n\n'):
        doc.add_paragraph(para.strip())

    # Add a short per-section explanation for quick reference
    doc.add_heading('Quick reference: commands', level=2)
    doc.add_paragraph('Install dependencies: python -m pip install -r requirements.txt')
    doc.add_paragraph('Dry-run (preview): python email_attachments_to_dropbox.py --dry-run')
    doc.add_paragraph('Run live: python email_attachments_to_dropbox.py')

    out_path = 'HSBC_script_documentation.docx'
    doc.save(out_path)
    print(f'Created documentation: {out_path}')

if __name__ == '__main__':
    make_doc()
